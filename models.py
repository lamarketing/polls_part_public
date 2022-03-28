import datetime
import os
from pathlib import Path
from shutil import make_archive
from typing import List, Union

from django.db import models
from django.db.models import QuerySet
from django.db.models.signals import post_delete
from docx import Document
from docxtpl import DocxTemplate
from docxcompose.composer import Composer

from consultants.models import Consultants


class Polls(models.Model):
    name = models.CharField('Название', max_length=400)
    description = models.TextField('Описание', max_length=10000, default="", blank=True)
    date_start = models.DateField('Дата старта', default=datetime.date.today)
    date_end = models.DateField('Дата окончания', default=datetime.date.today() + datetime.timedelta(7))

    @staticmethod
    def questions_all(poll) -> List[dict]:
        """
        Отдает список всех вопросов
        --- для конкретного голосования
        --- со всеми ответами для каждого вопроса
        """
        return [{'question': {'id': que.id, 'name': que.name},
                 'answers': [{'id': ans.id, 'name': ans.name}
                             for ans in que.answers.all()]}
                for que in poll.questions.all()]

    def __str__(self):
        return f'{self.id}: {self.name}'

    class Meta:
        verbose_name_plural = "1 Голосования"
        verbose_name = 'Голосование'
        ordering = ['-id']


class Answers(models.Model):
    name = models.CharField('Ответ', max_length=600)

    def __str__(self):
        return f'{self.name}'

    class Meta:
        verbose_name_plural = "3 Ответы"
        verbose_name = 'Ответ'


class Questions(models.Model):
    name = models.TextField('Вопрос', max_length=2000)
    poll = models.ForeignKey(Polls, on_delete=models.CASCADE, related_name='questions', verbose_name='Голосование')
    file = models.FileField(upload_to='library/', default="", blank=True, null=True)
    answers = models.ManyToManyField(Answers, verbose_name='Ответы')

    def __str__(self):
        return f'{self.id} - {self.poll.id}'

    class Meta:
        verbose_name_plural = "2 Вопросы"
        verbose_name = 'Вопрос'


class PollsResults(models.Model):

    def path_to_word(self, instance, filename):
        return f'consultants/polls/{instance.poll.id}/{instance.consultant.id}_{instance.id}.docx'

    consultant = models.ForeignKey(Consultants, on_delete=models.SET_NULL, null=True,
                                   related_name='polls_results', verbose_name='консультант')
    poll = models.ForeignKey(Polls, on_delete=models.CASCADE, related_name='polls_results', verbose_name='Голосование')
    result = models.JSONField('Результат')
    file = models.FileField('Word файл', null=True, blank=True)

    created = models.DateField(auto_now_add=True, auto_now=False)

    def __str__(self):
        return f'{self.consultant} - {self.id}'

    class Meta:
        verbose_name_plural = "4 Результаты голосований"
        verbose_name = 'Результат голосования'

    def save(self, *args, **kwargs):
        """
        Удаляем ворд файл результата, когда очищают файл в результате
        """
        if self.pk:
            old = PollsResults.objects.get(pk=self.pk)
            if old.file and not self.file:
                os.remove(f'media/{old.file}')
        super(PollsResults, self).save()

    @staticmethod
    def post_delete(sender, instance, **kwargs) -> None:
        """
        Удаляем ворд файл результата, когда удаляют результат
        """
        if instance.file:
            os.remove(f'media/{instance.file}')

    @staticmethod
    def update_statistics_from_one_result(questions: List[dict],
                                          statistics: List[dict],
                                          answers: List[dict],
                                          liter: Union[int, str] = 1) -> List[dict]:
        """
        Изменяет список "Вопрос-ответы" на основе ответов одного результата
        liter - значение(или метка "X"), которое добавится для выбранного ответа,
        для общей статистика +1, для одного результата "X"(можно другое)
        """
        for result_que in questions:
            # один вопрос из результата
            for Q_index, que in enumerate(answers):
                if que['question']['id'] == result_que['question']['id']:
                    # есть совпадение вопроса
                    for A_index, ans in enumerate(que['answers']):
                        # один ответ из ответов голосования
                        if ans['id'] == result_que['answer']['id']:
                            statistics[Q_index]['a'][A_index]['sum'] = (statistics[Q_index]['a'][A_index]['sum'] + 1) \
                                if liter == 1 else liter
        return statistics
    @property
    def create_word(self):
        doc = DocxTemplate("consultants/templates/template_poll_result.docx")
        context = self.result
        context['consultant_name'] = self.consultant.name
        context['consultant_last_name'] = self.consultant.last_name
        context['consultant_middle_name'] = self.consultant.middle_name
        context['consultant_attestat_numder'] = self.consultant.attestat_number
        context['consultant_date_end'] = self.consultant.date_end.strftime('%d.%m.%y') if self.consultant.date_end else "---"
        context['poll_date_start'] = self.poll.date_start.strftime('%d.%m.%y')
        context['poll_date_end'] = self.poll.date_end.strftime('%d.%m.%y')
        context['created'] = self.created.strftime('%d.%m.%y')

        answers: List[dict] = Polls.questions_all(self.poll)
        statistics: List[dict] = PollsResults.update_statistics_from_one_result(
            self.result['questions'],
            PollsStatistics.create_statistics(answers, '-'),
            answers,
            'X'
        )
        q_text = "\n\n--------------------------------------\n"\
            .join(["\n\n"
                  .join([q['q'], "\n"
                        .join([" "
                              .join([a['name'], str(a['sum'])])
                               for a in q['a']])])
                   for q in statistics])

        context['q_text'] = q_text
        doc.render(context)

        try:
            os.mkdir(f'media/polls/{self.poll.id}')
        except:
            pass
        path = f"media/polls/{self.poll.id}/{self.id}_{self.consultant.id}_{self.poll.id}.docx"
        doc.save(path)
        self.file = path[5:]
        self.save()

        return path


post_delete.connect(PollsResults.post_delete, PollsResults)


class PollsStatistics(models.Model):
    """
    Статистика конкретного голосования
    с zip-файлом всех сформированных ранее результатов этого голосования(word файлы)
    """
    poll = models.ForeignKey(Polls, on_delete=models.CASCADE, related_name='polls_statistics',
                             verbose_name='Голосование')
    statistics = models.TextField("Cтатистика", max_length=10000, null=True, blank=True)
    isFull = models.BooleanField('Полная', default=True,
                                 help_text="Голосование закончено и учтены все результаты голосов")
    zipFile = models.FileField('ZIP файл всех результатов', null=True, blank=True,
                               upload_to='polls/ZIP')
    wordBigFile = models.FileField('Один большой word', null=True, blank=True,
                               upload_to='polls/ZIP')

    created = models.DateField('Создан', auto_now_add=True, auto_now=False)

    @staticmethod
    def create_statistics(answers: List[dict],
                          liter: Union[int, str] = 0) -> List[dict]:
        """
        Создает список "Вопрос-ответы" для дальнейшего тестового отображения
        на основе всех вопросов со всеми ответами для конкретного голосования,
        где все ответы либо 0, либо указанный liter
        """
        return [{'q': q['question']['name'], 'a': [{'name': a['name'], 'sum': 0 if liter == 0 else liter}
                                                   for a in q['answers']]}
                for q in answers]

    def save(self, *args, **kwargs):
        """
        Только для создания.

        1) Создаем zip-файл со всеми word-файлами результатов этого голосования

        ---- Если у результата нет word-файла, то его голос не учитывается

        ---- Если не все голоса являются учтенными или голосование еще не закончилось,
        то isFool = False (чисто для восприятия)

        2) Все word фалы в один для печати

        3) Если архив создан, то формируем статистику для конкретного голосования
        """
        if not self.pk:
            poll: Polls = self.poll
            # Проверяем есть папка с результатами
            path_dir = f'media/polls/{poll.id}'
            is_dir: bool = os.path.isdir(path_dir)
            if is_dir:
                try:
                    archive = make_archive(f'media/polls/ZIP/{poll.id}', "zip", path_dir)
                except:
                    archive = None

                if archive:
                    # # ВСЕ word файлы в один файл
                    # init_doc = Composer(Document())
                    # empty = Document('media/polls/ZIP/empty_dont_delete.docx')
                    # for doc in Path(path_dir).glob('*.docx'):
                    #     init_doc.append(Document(doc))
                    #     init_doc.append(empty)
                    # init_doc.save(f'media/polls/ZIP/{poll.id}.docx')

                    # СТАТИСТИКА
                    results: QuerySet[PollsResults] = PollsResults.objects.filter(poll=self.poll)
                    results_lens = len(results)
                    results_archive_lens = 0
                    answers: List[dict] = Polls.questions_all(poll)
                    statistics: List[dict] = self.create_statistics(answers)
                    for result in results:
                        # один результат
                        if result.file:
                            # если есть файл ворд
                            results_archive_lens += 1
                            statistics = PollsResults.update_statistics_from_one_result(
                                result.result['questions'], statistics, answers)

                    if statistics and results_archive_lens:
                        q_text = "\n\n"\
                            .join(["\n\n"
                                  .join([q['q'], "\n"
                                        .join([" "
                                              .join([a['name'], str(a['sum']),
                                                     '(', str(int(a['sum'] * 100/results_archive_lens)), '% )'])
                                               for a in q['a']])])
                                   for q in statistics])
                        stat_text = f"""
                        Всего проголосовавших: {results_lens} 
                        Учтенные голоса: {results_archive_lens}
                        
                        {q_text}
                        """
                        self.statistics = stat_text
                        today = datetime.date.today()
                        if results_lens > results_archive_lens or poll.date_end >= today:
                            self.isFull = False
                        self.zipFile = "/".join(archive.split('/')[-3:])
                        self.wordBigFile = ''
                        super().save(*args, **kwargs)

    def __str__(self):
        return f'{self.id} - голосование({self.id})'

    class Meta:
        verbose_name_plural = "5 Статистики голосований"
        verbose_name = 'Статистика голосования'


class PollsMailing(models.Model):
    consultants_add = models.ManyToManyField(Consultants, related_name='polls_mailing', verbose_name='Отправили', blank=True)
    consultants_error = models.ManyToManyField(Consultants, verbose_name='Ошибки', blank=True)
    errors = models.TextField('Общий текст ошибок', max_length=200000, blank=True, null=True)
    slug = models.CharField('Slug', max_length=20, default="", null=True, unique=True)

    created = models.DateField('Создан', auto_now_add=True, auto_now=False)
