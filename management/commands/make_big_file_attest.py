import os
from datetime import date, datetime
from pathlib import Path
from random import randint
from typing import Optional, List

from django.core.management import BaseCommand, CommandError
from django.db.models import QuerySet
from docx import Document
from docxcompose.composer import Composer

from consultants.models import Consultants
from polls.models import PollsStatistics, PollsResults, Polls


class Command(BaseCommand):
    """
    1) Найти результаты голосов
    2) Собрать файлы
    3) Сделать один файл
    """
    help = 'Один большой файл только для аттестованных'

    def add_arguments(self, parser):
        parser.add_argument('poll_id', nargs='+', type=int)

    def handle(self, *args, **options):
        if not options['poll_id']:
            raise CommandError('Укажите poll_id')

        poll_id = options['poll_id'][0]
        prs = PollsResults.objects.filter(
            poll__id=poll_id,
        ).exclude(
            consultant__attestat_number__isnull=False,
            consultant__date_end__gt=datetime(
                year=2022, month=2, day=12, hour=23, minute=0, second=0, microsecond=0
            ).astimezone(),
        ).values_list('id', flat=True)
        print(len(prs))
        path_dir = f'media/polls/{poll_id}'
        init_docs = [
            Composer(Document()),
            Composer(Document()),
            Composer(Document()),
            Composer(Document()),
            Composer(Document()),
            Composer(Document())]
        empty = Document('media/polls/ZIP/empty_dont_delete.docx')
        j = 1
        for doc in Path(path_dir).glob('*.docx'):
            pr_id = int(doc.name.split("_")[0])
            if pr_id in prs:
                print(j)
                if j <= 500:
                    init_doc = init_docs[0]
                elif j <= 1000:
                    init_doc = init_docs[1]
                elif j <= 1500:
                    init_doc = init_docs[2]
                elif j <= 2000:
                    init_doc = init_docs[3]
                elif j <= 2500:
                    init_doc = init_docs[4]
                else:
                    init_doc = init_docs[5]

                init_doc.append(Document(doc))
                init_doc.append(empty)
                j += 1

        for index, init_d in enumerate(init_docs, start=1):
            init_d.save(f'media/polls/ZIP/NA_{index * 500}.docx')
