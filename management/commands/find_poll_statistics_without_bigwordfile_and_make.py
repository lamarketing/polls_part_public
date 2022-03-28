import os
from pathlib import Path
from typing import Optional

from django.core.management import BaseCommand
from django.db.models import QuerySet
from docx import Document
from docxcompose.composer import Composer

from polls.models import PollsStatistics



class Command(BaseCommand):
    """
    Берет последнюю статистику голосования без большого word файла
    Создает большой word файл и сохраняет его в поле статистики
    """
    help = 'Изменение статуса консультанта за полгода до окончания'

    def handle(self, *args, **options):
        poll: Optional[PollsStatistics] = PollsStatistics.objects.filter(
            wordBigFile=''
        ).last()
        path_dir = f'media/polls/3'
        is_dir: bool = os.path.isdir(path_dir)
        if is_dir and poll:
            # ВСЕ word файлы в один файл
            init_doc = Composer(Document())
            empty = Document('media/polls/ZIP/empty_dont_delete.docx')
            for doc in Path(path_dir).glob('*.docx'):
                init_doc.append(Document(doc))
                init_doc.append(empty)
            init_doc.save(f'media/polls/ZIP/{poll.id}.docx')

            poll.wordBigFile = f'media/polls/ZIP/{poll.id}.docx'
            poll.save()
